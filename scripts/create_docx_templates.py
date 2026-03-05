"""Script to create .docx templates for BIM proposals and quotes.

Run: python scripts/create_docx_templates.py
Generates:
  - worker/templates/documents/propuesta_bim.docx
  - worker/templates/documents/cotizacion_bim.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "worker", "templates", "documents")

BRAND_BLUE = RGBColor(0x1A, 0x47, 0x7A)
BRAND_GRAY = RGBColor(0x66, 0x66, 0x66)
BRAND_LIGHT = RGBColor(0x99, 0x99, 0x99)
BRAND_DARK = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "1A477A"
TABLE_ALT_BG = "F2F6FA"


def _set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading_el = tc_pr.makeelement(
        qn("w:shd"),
        {qn("w:fill"): hex_color, qn("w:val"): "clear"},
    )
    tc_pr.append(shading_el)


def _add_styled_table(doc, headers: list, rows: list, col_widths=None):
    """Create a styled table with header row and optional data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        _set_cell_bg(cell, TABLE_HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            if r_idx % 2 == 1:
                _set_cell_bg(cell, TABLE_ALT_BG)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def _add_heading_styled(doc, text: str, level: int = 1):
    """Add a heading with brand color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_BLUE
    return h


def _add_footer(doc):
    """Add standard footer."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("David Moreira — Consultoria BIM & Tecnologia")
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("contacto@umbral.dev | umbral.dev")
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_LIGHT


def create_propuesta_bim():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ── 1. Portada ──
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Logo Empresa]")
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_LIGHT

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROPUESTA DE CONSULTORIA BIM")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BRAND_BLUE

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ proyecto }}")
    run.font.size = Pt(18)
    run.font.color.rgb = BRAND_DARK

    doc.add_paragraph()

    cover_info = [
        ("Cliente:", "{{ cliente }}"),
        ("Fecha:", "{{ fecha }}"),
        ("Version:", "{{ version }}"),
        ("Referencia:", "{{ referencia }}"),
    ]
    for label, value in cover_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label} ")
        run.font.size = Pt(12)
        run.font.color.rgb = BRAND_GRAY
        run = p.add_run(value)
        run.font.size = Pt(12)

    doc.add_page_break()

    # ── 2. Resumen Ejecutivo ──
    _add_heading_styled(doc, "1. Resumen Ejecutivo")

    doc.add_paragraph("{{ resumen_ejecutivo }}")
    doc.add_paragraph()
    doc.add_paragraph(
        "Esta propuesta detalla los servicios de consultoria BIM requeridos "
        "para el proyecto {{ proyecto }}, incluyendo alcance, metodologia, "
        "equipo, cronograma e inversion."
    )
    doc.add_paragraph()

    # ── 3. Alcance de Servicios BIM ──
    _add_heading_styled(doc, "2. Alcance de Servicios BIM")

    doc.add_paragraph("{{ alcance_general }}")
    doc.add_paragraph()

    _add_heading_styled(doc, "2.1 Niveles de Desarrollo (LOD)", level=2)

    _add_styled_table(
        doc,
        headers=["LOD", "Fase", "Contenido del modelo", "Usos BIM"],
        rows=[
            [
                "LOD 200",
                "Diseno esquematico",
                "Elementos genericos con dimensiones aproximadas, "
                "materiales indicativos",
                "Coordinacion espacial, estimacion de costos preliminar",
            ],
            [
                "LOD 300",
                "Desarrollo de diseno",
                "Elementos con dimensiones precisas, materiales "
                "definidos, propiedades tecnicas",
                "Coordinacion 3D, clash detection, quantity takeoff",
            ],
            [
                "LOD 400",
                "Documentacion de construccion",
                "Detalle constructivo, conexiones, fabricacion, "
                "secuencia de montaje",
                "Planificacion 4D, prefabricacion, control de obra",
            ],
        ],
        col_widths=[2, 3, 6, 5],
    )
    doc.add_paragraph()

    # ── 4. Metodologia ──
    _add_heading_styled(doc, "3. Metodologia")

    doc.add_paragraph("{{ metodologia }}")
    doc.add_paragraph()

    _add_heading_styled(doc, "3.1 Entregables por fase", level=2)

    _add_styled_table(
        doc,
        headers=["Fase", "Entregables", "Formato", "Criterio de aceptacion"],
        rows=[
            [
                "Diseno",
                "Modelo BIM federado, planos coordinados, "
                "reporte de interferencias",
                "RVT / IFC / NWD / PDF",
                "Cero clashes criticos entre disciplinas",
            ],
            [
                "Construccion",
                "Modelo 4D/5D, planificacion de montaje, "
                "quantity takeoff, reportes de avance",
                "NWD / MPP / XLSX / PDF",
                "Modelo actualizado semanalmente, "
                "reporte de avance mensual",
            ],
            [
                "Operacion",
                "Modelo As-Built, manual de operacion BIM, "
                "datos COBie para facility management",
                "IFC / COBie / PDF",
                "Entrega conforme a requerimientos del mandante",
            ],
        ],
        col_widths=[3, 5, 4, 5],
    )
    doc.add_paragraph()

    # ── 5. Equipo Propuesto ──
    _add_heading_styled(doc, "4. Equipo Propuesto")

    _add_styled_table(
        doc,
        headers=["Rol", "Perfil", "Dedicacion", "Horas estimadas"],
        rows=[
            [
                "BIM Manager / Consultor Principal",
                "Arquitecto con +10 anios en BIM. Certificado buildingSMART "
                "y Autodesk. Experiencia en coordinacion multidisciplinaria",
                "{{ dedicacion_bim_manager }}",
                "{{ horas_bim_manager }}",
            ],
            [
                "Coordinador BIM",
                "Especialista en clash detection, federacion de modelos "
                "y gestion de entregables BIM",
                "{{ dedicacion_coordinador }}",
                "{{ horas_coordinador }}",
            ],
            [
                "Modelador BIM",
                "Tecnico con dominio de Revit, Navisworks y herramientas "
                "de documentacion",
                "{{ dedicacion_modelador }}",
                "{{ horas_modelador }}",
            ],
        ],
        col_widths=[4, 6, 3, 3],
    )
    doc.add_paragraph()

    # ── 6. Cronograma ──
    _add_heading_styled(doc, "5. Cronograma")

    _add_styled_table(
        doc,
        headers=["Etapa", "Actividades clave", "Duracion", "Hito de cierre"],
        rows=[
            [
                "1. Kick-off y diagnostico",
                "Reunion de inicio, revision de documentacion existente, "
                "definicion de BEP (BIM Execution Plan)",
                "{{ duracion_etapa_1 }}",
                "BEP aprobado por el cliente",
            ],
            [
                "2. Modelado y coordinacion",
                "Modelado por disciplina, coordinacion 3D, "
                "clash detection, reportes semanales",
                "{{ duracion_etapa_2 }}",
                "Modelo federado sin clashes criticos",
            ],
            [
                "3. Documentacion",
                "Generacion de planos, listas de cantidades, "
                "especificaciones tecnicas",
                "{{ duracion_etapa_3 }}",
                "Set de planos aprobado",
            ],
            [
                "4. Entrega y cierre",
                "Modelo As-Built, capacitacion al equipo del cliente, "
                "transferencia de conocimiento",
                "{{ duracion_etapa_4 }}",
                "Acta de cierre firmada",
            ],
        ],
        col_widths=[3.5, 5.5, 3, 4],
    )
    doc.add_paragraph()

    # ── 7. Honorarios ──
    _add_heading_styled(doc, "6. Inversion")

    _add_styled_table(
        doc,
        headers=["Item", "Descripcion", "Cantidad", "Valor unitario", "Subtotal"],
        rows=[
            [
                "1",
                "Consultoria BIM Manager",
                "{{ cantidad_bim_manager }}",
                "{{ valor_unit_bim_manager }}",
                "{{ subtotal_bim_manager }}",
            ],
            [
                "2",
                "Coordinacion BIM",
                "{{ cantidad_coordinador }}",
                "{{ valor_unit_coordinador }}",
                "{{ subtotal_coordinador }}",
            ],
            [
                "3",
                "Modelado BIM",
                "{{ cantidad_modelador }}",
                "{{ valor_unit_modelador }}",
                "{{ subtotal_modelador }}",
            ],
            [
                "4",
                "Licencias y software",
                "{{ cantidad_licencias }}",
                "{{ valor_unit_licencias }}",
                "{{ subtotal_licencias }}",
            ],
        ],
        col_widths=[1.5, 5, 2.5, 3, 3],
    )
    doc.add_paragraph()

    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    totals_data = [
        ("Subtotal:", "{{ subtotal }}"),
        ("IVA (19%):", "{{ iva }}"),
        ("TOTAL:", "{{ total }}"),
    ]
    for i, (label, value) in enumerate(totals_data):
        totals_table.rows[i].cells[0].text = label
        totals_table.rows[i].cells[1].text = value
        for cell in totals_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Calibri"
                    if i == 2:
                        run.bold = True

    doc.add_paragraph()

    # ── 8. Condiciones Generales ──
    _add_heading_styled(doc, "7. Condiciones Generales")

    conditions = [
        "Validez de la propuesta: 30 dias corridos desde la fecha de emision.",
        "Forma de pago: {{ condiciones_pago }}",
        "Los plazos indicados comienzan a partir de la firma del contrato "
        "y la recepcion de la documentacion base por parte del cliente.",
        "El cliente debe designar un interlocutor unico para la coordinacion "
        "del proyecto.",
        "Cualquier modificacion al alcance sera evaluada y cotizada por separado.",
        "Los modelos BIM se entregan en formato nativo (RVT) e IFC. "
        "Los planos en PDF y DWG.",
    ]
    for cond in conditions:
        doc.add_paragraph(cond, style="List Bullet")

    doc.add_paragraph()

    # ── 9. Firma y Aceptacion ──
    _add_heading_styled(doc, "8. Firma y Aceptacion")

    doc.add_paragraph(
        "Con la firma de este documento, ambas partes aceptan los terminos "
        "y condiciones descritos en esta propuesta."
    )
    doc.add_paragraph()

    sign_table = doc.add_table(rows=4, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.rows[0].cells[0].text = "Por el Consultor:"
    sign_table.rows[0].cells[1].text = "Por el Cliente:"
    sign_table.rows[1].cells[0].text = ""
    sign_table.rows[1].cells[1].text = ""
    sign_table.rows[2].cells[0].text = "________________________"
    sign_table.rows[2].cells[1].text = "________________________"
    sign_table.rows[3].cells[0].text = "David Moreira Mercado"
    sign_table.rows[3].cells[1].text = "{{ nombre_representante }}"

    for row in sign_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Calibri"

    _add_footer(doc)

    out = os.path.join(TEMPLATES_DIR, "propuesta_bim.docx")
    doc.save(out)
    print(f"Created: {out}")


def create_cotizacion_bim():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ── 1. Encabezado empresa ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("David Moreira — Consultoria BIM & Tecnologia")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BRAND_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("contacto@umbral.dev | umbral.dev | Santiago, Chile")
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("COTIZACION")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BRAND_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Ref: {{ referencia }}")
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND_GRAY

    doc.add_paragraph()

    # ── 2. Datos del cliente ──
    _add_heading_styled(doc, "Datos del Cliente", level=2)

    client_table = doc.add_table(rows=5, cols=2)
    client_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    client_data = [
        ("Empresa:", "{{ empresa_cliente }}"),
        ("Contacto:", "{{ contacto_cliente }}"),
        ("Proyecto:", "{{ proyecto }}"),
        ("Fecha:", "{{ fecha }}"),
        ("Validez:", "{{ vigencia }}"),
    ]
    for i, (label, value) in enumerate(client_data):
        row = client_table.rows[i]
        row.cells[0].text = ""
        p = row.cells[0].paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        row.cells[1].text = ""
        p = row.cells[1].paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    doc.add_paragraph()

    # ── 3. Objeto de la cotizacion ──
    _add_heading_styled(doc, "Objeto de la Cotizacion", level=2)
    doc.add_paragraph("{{ objeto_cotizacion }}")
    doc.add_paragraph()

    # ── 4. Detalle de items ──
    _add_heading_styled(doc, "Detalle de Servicios")

    _add_styled_table(
        doc,
        headers=["Item", "Descripcion", "Hrs", "Valor Unit.", "Total"],
        rows=[
            [
                "1",
                "Diagnostico BIM: relevamiento de procesos actuales, "
                "identificacion de brechas, recomendaciones",
                "{{ hrs_diagnostico }}",
                "{{ valor_diagnostico }}",
                "{{ total_diagnostico }}",
            ],
            [
                "2",
                "BIM Execution Plan (BEP): definicion de usos BIM, "
                "roles, entregables, matriz de responsabilidades",
                "{{ hrs_bep }}",
                "{{ valor_bep }}",
                "{{ total_bep }}",
            ],
            [
                "3",
                "Coordinacion 3D: federacion de modelos, clash detection, "
                "reportes semanales de interferencias",
                "{{ hrs_coordinacion }}",
                "{{ valor_coordinacion }}",
                "{{ total_coordinacion }}",
            ],
            [
                "4",
                "Capacitacion: workshops de Revit, Navisworks y flujos BIM "
                "para el equipo del proyecto",
                "{{ hrs_capacitacion }}",
                "{{ valor_capacitacion }}",
                "{{ total_capacitacion }}",
            ],
            [
                "5",
                "Soporte continuo: asesoria remota, revision de modelos, "
                "resolucion de consultas tecnicas",
                "{{ hrs_soporte }}",
                "{{ valor_soporte }}",
                "{{ total_soporte }}",
            ],
        ],
        col_widths=[1.2, 7, 1.5, 2.5, 2.5],
    )

    doc.add_paragraph()

    # ── 5. Subtotal / IVA / Total ──
    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    totals_data = [
        ("Subtotal:", "{{ subtotal }}"),
        ("IVA (19%):", "{{ iva }}"),
        ("TOTAL:", "{{ total }}"),
    ]
    for i, (label, value) in enumerate(totals_data):
        row = totals_table.rows[i]
        row.cells[0].text = ""
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p0.add_run(label)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i == 2:
            run.bold = True

        row.cells[1].text = ""
        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p1.add_run(value)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i == 2:
            run.bold = True
            run.font.color.rgb = BRAND_BLUE

    doc.add_paragraph()

    # ── 6. Validez y condiciones de pago ──
    _add_heading_styled(doc, "Condiciones", level=2)

    conditions = [
        "Validez de la cotizacion: {{ vigencia }} desde la fecha de emision.",
        "Forma de pago: {{ condiciones_pago }}",
        "Los valores estan expresados en {{ moneda }}.",
        "Plazo estimado de ejecucion: {{ plazo_ejecucion }}.",
        "Esta cotizacion no incluye: {{ exclusiones }}.",
        "Los precios pueden ajustarse si el alcance cambia respecto "
        "a lo descrito en esta cotizacion.",
    ]
    for cond in conditions:
        doc.add_paragraph(cond, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph("{{ notas_adicionales }}")

    doc.add_paragraph()

    # ── 7. Firma ──
    _add_heading_styled(doc, "Aceptacion", level=2)

    doc.add_paragraph(
        "Acepto los terminos y condiciones de esta cotizacion:"
    )
    doc.add_paragraph()

    sign_table = doc.add_table(rows=3, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.rows[0].cells[0].text = "________________________"
    sign_table.rows[0].cells[1].text = "________________________"
    sign_table.rows[1].cells[0].text = "David Moreira Mercado"
    sign_table.rows[1].cells[1].text = "{{ nombre_cliente }}"
    sign_table.rows[2].cells[0].text = "Consultor BIM"
    sign_table.rows[2].cells[1].text = "{{ cargo_cliente }}"

    for row in sign_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

    _add_footer(doc)

    out = os.path.join(TEMPLATES_DIR, "cotizacion_bim.docx")
    doc.save(out)
    print(f"Created: {out}")


if __name__ == "__main__":
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    create_propuesta_bim()
    create_cotizacion_bim()
    print("Done!")
