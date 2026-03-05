"""Tests for BIM skills (IFC, Speckle, BIM Coordination) and DOCX templates."""

import os
import re

import pytest
import yaml
from docx import Document

SKILLS_DIR = os.path.join(
    os.path.dirname(__file__),
    "..",
    "openclaw",
    "workspace-templates",
    "skills",
)
TEMPLATES_DIR = os.path.join(
    os.path.dirname(__file__),
    "..",
    "worker",
    "templates",
    "documents",
)


def _load_skill_frontmatter(skill_name: str) -> dict:
    """Parse YAML frontmatter from a SKILL.md file."""
    path = os.path.join(SKILLS_DIR, skill_name, "SKILL.md")
    with open(path, encoding="utf-8") as f:
        content = f.read()
    match = re.match(r"^---\n(.*?)\n---", content, re.DOTALL)
    assert match, f"No YAML frontmatter found in {path}"
    return yaml.safe_load(match.group(1))


# ── IFC Skill ──


def test_ifc_skill_has_required_frontmatter():
    fm = _load_skill_frontmatter("ifc-python")
    assert fm["name"] == "ifc-python"
    assert "description" in fm
    assert len(fm["description"]) > 20
    assert "metadata" in fm
    assert "openclaw" in fm["metadata"]
    assert "requires" in fm["metadata"]["openclaw"]
    assert "env" in fm["metadata"]["openclaw"]["requires"]
    assert fm["metadata"]["openclaw"]["requires"]["env"] == []


# ── Speckle Skill ──


def test_speckle_skill_has_speckle_token_env():
    fm = _load_skill_frontmatter("speckle-dalux-powerbi")
    env_vars = fm["metadata"]["openclaw"]["requires"]["env"]
    assert "SPECKLE_TOKEN" in env_vars


# ── BIM Coordination Skill ──


def test_bim_coordination_skill_no_env_required():
    fm = _load_skill_frontmatter("bim-coordination")
    env_vars = fm["metadata"]["openclaw"]["requires"]["env"]
    assert env_vars == []


# ── Common validations ──


@pytest.mark.parametrize(
    "skill_name",
    ["ifc-python", "speckle-dalux-powerbi", "bim-coordination"],
)
def test_all_new_skills_have_description(skill_name):
    fm = _load_skill_frontmatter(skill_name)
    assert "description" in fm
    assert len(fm["description"]) > 20


@pytest.mark.parametrize(
    "skill_name",
    ["ifc-python", "speckle-dalux-powerbi", "bim-coordination"],
)
def test_all_new_skills_have_emoji(skill_name):
    fm = _load_skill_frontmatter(skill_name)
    emoji = fm["metadata"]["openclaw"]["emoji"]
    assert emoji is not None
    assert len(emoji) > 0


# ── Propuesta template ──


def test_propuesta_template_exists_and_has_content():
    path = os.path.join(TEMPLATES_DIR, "propuesta_bim.docx")
    assert os.path.exists(path), "propuesta_bim.docx does not exist"
    size = os.path.getsize(path)
    assert size > 10_000, f"propuesta_bim.docx too small ({size} bytes)"


def test_propuesta_template_has_sections():
    path = os.path.join(TEMPLATES_DIR, "propuesta_bim.docx")
    doc = Document(path)

    headings = [
        p.text.strip()
        for p in doc.paragraphs
        if p.style.name.startswith("Heading")
    ]

    expected_keywords = [
        "Resumen Ejecutivo",
        "Alcance",
        "Metodologia",
        "Equipo",
        "Cronograma",
        "Inversion",
        "Condiciones",
        "Firma",
    ]

    for keyword in expected_keywords:
        found = any(keyword.lower() in h.lower() for h in headings)
        assert found, (
            f"Expected heading containing '{keyword}' not found. "
            f"Headings: {headings}"
        )


# ── Cotizacion template ──


def test_cotizacion_template_exists_and_has_content():
    path = os.path.join(TEMPLATES_DIR, "cotizacion_bim.docx")
    assert os.path.exists(path), "cotizacion_bim.docx does not exist"
    size = os.path.getsize(path)
    assert size > 10_000, f"cotizacion_bim.docx too small ({size} bytes)"


def test_cotizacion_template_has_table():
    path = os.path.join(TEMPLATES_DIR, "cotizacion_bim.docx")
    doc = Document(path)

    assert len(doc.tables) >= 1, "cotizacion_bim.docx has no tables"

    service_table = None
    for table in doc.tables:
        header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "item" in header_texts and "total" in header_texts:
            service_table = table
            break

    assert service_table is not None, (
        "No service detail table with 'Item' and 'Total' headers found. "
        f"Tables found: {len(doc.tables)}"
    )
    assert len(service_table.rows) >= 3, (
        "Service table should have header + at least 2 data rows"
    )
